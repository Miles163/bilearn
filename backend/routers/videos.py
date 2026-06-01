import json
import io
import os
from fastapi import APIRouter, Depends, HTTPException, Header
from fastapi.responses import PlainTextResponse, Response, FileResponse
from sqlalchemy.orm import Session

from database import get_db, SessionLocal
from models import Video, Note, ReviewCard, ReviewLog
from schemas import VideoImport, VideoOut, NoteOut, GenerateIn
from services.bilibili import extract_bvid, get_video_info, get_subtitle
from services.youtube import extract_yt_id, get_video_info_and_subtitles
from services.llm_service import generate_notes
from services.card_gen import create_cards_from_notes

router = APIRouter(prefix="/api/videos", tags=["videos"])


def _detect_source(url: str) -> str:
    if "bilibili.com" in url:
        return "bilibili"
    elif "youtube.com" in url or "youtu.be" in url:
        return "youtube"
    raise ValueError("不支持的视频来源，请使用 B站 或 YouTube 链接")


@router.post("", response_model=VideoOut)
def import_video(data: VideoImport, db: Session = Depends(get_db), x_bili_token: str | None = Header(None)):
    try:
        source = _detect_source(data.url)
    except ValueError as e:
        raise HTTPException(400, str(e))

    yt_subs = []
    try:
        if source == "bilibili":
            bvid = extract_bvid(data.url)
            existing = db.query(Video).filter(Video.bvid == bvid).first()
            if existing:
                return existing
            info = get_video_info(bvid, x_bili_token)
            info["source"] = "bilibili"
        else:
            vid = extract_yt_id(data.url)
            existing = db.query(Video).filter(Video.bvid == vid).first()
            if existing:
                return existing
            result = get_video_info_and_subtitles(data.url)
            yt_subs = result.get("subtitles", [])
            info = {k: v for k, v in result.items() if k != "subtitles"}
            info["source"] = "youtube"
    except ValueError as e:
        raise HTTPException(400, str(e))

    video = Video(url=data.url, **info)
    db.add(video)
    db.commit()
    db.refresh(video)

    if source == "youtube" and yt_subs:
        subtitle_texts = []
        for s in yt_subs:
            subtitle_texts.append(f"[{s['lang']}] {s['text']}")
        note = Note(video_id=video.id, subtitle_text="\n\n".join(subtitle_texts))
        db.add(note)
        db.commit()

    return video


@router.get("", response_model=list[VideoOut])
def list_videos(db: Session = Depends(get_db)):
    return db.query(Video).order_by(Video.created_at.desc()).all()


@router.get("/{video_id}", response_model=VideoOut)
def get_video(video_id: int, db: Session = Depends(get_db)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")
    return v


@router.delete("/{video_id}")
def delete_video(video_id: int, db: Session = Depends(get_db)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")
    notes = db.query(Note).filter(Note.video_id == video_id).all()
    for note in notes:
        cards = db.query(ReviewCard).filter(ReviewCard.note_id == note.id).all()
        for card in cards:
            db.query(ReviewLog).filter(ReviewLog.card_id == card.id).delete()
            db.delete(card)
        db.delete(note)
    db.delete(v)
    db.commit()
    return {"ok": True}


@router.post("/{video_id}/generate", response_model=NoteOut)
def generate(video_id: int, data: GenerateIn | None = None, db: Session = Depends(get_db), x_bili_token: str | None = Header(None)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")

    note = db.query(Note).filter(Note.video_id == video_id).first()
    if note and note.summary and len(note.subtitle_text or "") >= 100:
        return note

    text = data.text if data and data.text else ""

    if not text:
        if v.source == "youtube":
            if note and note.subtitle_text:
                text = note.subtitle_text
            else:
                result = get_video_info_and_subtitles(v.url)
                subs = result.get("subtitles", [])
                text = "\n\n".join(f"[{s['lang']}] {s['text']}" for s in subs) if subs else ""
        else:
            text = get_subtitle(v.bvid, v.cid, x_bili_token)

    if not text:
        text = f"标题：{v.title}\n简介：{v.description}" if v.description != "-" else f"标题：{v.title}"

    subtitle = text.strip()
    if len(subtitle) < 100:
        raise HTTPException(400, "字幕过短（<100字），无法生成有效笔记。请先点击「从音频生成字幕」按钮获取完整字幕。")

    result, _usage = generate_notes(subtitle)

    if note is None:
        note = Note(video_id=video_id)
        db.add(note)

    note.summary = result.get("summary", "")
    note.key_points = json.dumps(result.get("key_points", []), ensure_ascii=False)
    note.subtitle_text = subtitle
    note.cleaned_subtitle = result.get("cleaned_subtitle", "")
    note.translated_subtitle = result.get("translated_subtitle", "")
    note.usage = json.dumps(_usage, ensure_ascii=False)
    db.commit()
    db.refresh(note)

    if result.get("cards"):
        create_cards_from_notes(note, result["cards"], db)

    return note


@router.get("/{video_id}/note", response_model=NoteOut | None)
def get_note(video_id: int, db: Session = Depends(get_db)):
    return db.query(Note).filter(Note.video_id == video_id).first()


# ── Async transcription with progress ──

import threading
import uuid

_transcribe_tasks: dict[str, dict] = {}


def _run_transcribe(video_id: int, task_id: str):
    from services.transcribe import transcribe_from_youtube, cleanup_temp
    from database import SessionLocal

    db = SessionLocal()
    try:
        v = db.query(Video).filter(Video.id == video_id).first()
        if not v:
            _transcribe_tasks[task_id] = {"status": "error", "error": "Video not found", "progress": 0}
            return

        _transcribe_tasks[task_id] = {"status": "downloading", "progress": 0, "phase": "下载音频..."}

        def on_progress(p: float):
            _transcribe_tasks[task_id] = {"status": "transcribing", "progress": p, "phase": f"转写中 {p*100:.0f}%"}

        _transcribe_tasks[task_id] = {"status": "transcribing", "progress": 0, "phase": "转写中 0%"}
        text = transcribe_from_youtube(v.url, v.bvid, on_progress=on_progress)

        cleanup_temp(v.bvid)

        if not text or len(text.strip()) <= 20:
            _transcribe_tasks[task_id] = {"status": "error", "error": "转录失败，视频可能没有音频轨道", "progress": 0}
            return

        _transcribe_tasks[task_id] = {"status": "generating", "progress": 0.95, "phase": "AI 生成笔记..."}

        subtitle = text.strip()
        result, usage = generate_notes(subtitle)

        note = db.query(Note).filter(Note.video_id == video_id).first()
        if note is None:
            note = Note(video_id=video_id)
            db.add(note)

        note.summary = result.get("summary", "")
        note.key_points = json.dumps(result.get("key_points", []), ensure_ascii=False)
        note.subtitle_text = subtitle
        note.cleaned_subtitle = result.get("cleaned_subtitle", "")
        note.translated_subtitle = result.get("translated_subtitle", "")
        note.usage = json.dumps(usage, ensure_ascii=False)
        db.commit()
        db.refresh(note)

        if result.get("cards"):
            from services.card_gen import create_cards_from_notes
            create_cards_from_notes(note, result["cards"], db)

        from schemas import NoteOut
        _transcribe_tasks[task_id] = {"status": "done", "progress": 1.0, "note": NoteOut.model_validate(note, from_attributes=True).model_dump()}
    except Exception as e:
        _transcribe_tasks[task_id] = {"status": "error", "error": str(e), "progress": 0}
    finally:
        db.close()


@router.post("/{video_id}/transcribe")
def start_transcribe(video_id: int, db: Session = Depends(get_db)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")

    note = db.query(Note).filter(Note.video_id == video_id).first()
    if note and note.subtitle_text and len(note.subtitle_text) > 50:
        if not note.subtitle_text.startswith("标题："):
            return {"task_id": None, "note": NoteOut.model_validate(note, from_attributes=True).model_dump()}

    task_id = str(uuid.uuid4())
    _transcribe_tasks[task_id] = {"status": "queued", "progress": 0, "phase": "等待中..."}

    t = threading.Thread(target=_run_transcribe, args=(video_id, task_id), daemon=True)
    t.start()

    return {"task_id": task_id, "note": None}


@router.get("/{video_id}/transcribe/status")
def get_transcribe_status(video_id: int, task_id: str, db: Session = Depends(get_db)):
    task = _transcribe_tasks.get(task_id)
    if not task:
        return {"status": "not_found", "progress": 0}

    if task["status"] == "done":
        return task
    return {"status": task["status"], "progress": task["progress"], "phase": task.get("phase", ""), "error": task.get("error", "")}


_generate_tasks: dict[str, dict] = {}


def _run_generate(video_id: int, task_id: str, text_override: str, token: str | None):
    from database import SessionLocal

    db = SessionLocal()
    try:
        _generate_tasks[task_id] = {"status": "running", "progress": 0.1, "phase": "准备字幕..."}

        v = db.query(Video).filter(Video.id == video_id).first()
        if not v:
            _generate_tasks[task_id] = {"status": "error", "error": "Video not found"}
            return

        text = text_override
        if not text:
            note = db.query(Note).filter(Note.video_id == video_id).first()
            _generate_tasks[task_id] = {"status": "running", "progress": 0.2, "phase": "获取字幕..."}

            if v.source == "youtube":
                if note and note.subtitle_text:
                    text = note.subtitle_text
                else:
                    from services.youtube import get_video_info_and_subtitles
                    result = get_video_info_and_subtitles(v.url)
                    subs = result.get("subtitles", [])
                    text = "\n\n".join(f"[{s['lang']}] {s['text']}" for s in subs) if subs else ""
            else:
                text = get_subtitle(v.bvid, v.cid, token)

        if not text:
            text = f"标题：{v.title}\n简介：{v.description}" if v.description != "-" else f"标题：{v.title}"

        subtitle = text.strip()
        if len(subtitle) < 100:
            _generate_tasks[task_id] = {"status": "error", "error": "字幕过短（<100字），无法生成有效笔记。请先点击「从音频生成字幕」按钮获取完整字幕。", "progress": 0}
            return

        _generate_tasks[task_id] = {"status": "running", "progress": 0.4, "phase": "AI 分析中..."}
        result, usage = generate_notes(subtitle)

        note = db.query(Note).filter(Note.video_id == video_id).first()
        if note is None:
            note = Note(video_id=video_id)
            db.add(note)

        _generate_tasks[task_id] = {"status": "running", "progress": 0.8, "phase": "保存笔记..."}

        note.summary = result.get("summary", "")
        note.key_points = json.dumps(result.get("key_points", []), ensure_ascii=False)
        note.subtitle_text = subtitle
        note.cleaned_subtitle = result.get("cleaned_subtitle", "")
        note.translated_subtitle = result.get("translated_subtitle", "")
        note.usage = json.dumps(usage, ensure_ascii=False)
        db.commit()
        db.refresh(note)

        if result.get("cards"):
            from services.card_gen import create_cards_from_notes
            create_cards_from_notes(note, result["cards"], db)

        _generate_tasks[task_id] = {"status": "running", "progress": 0.95, "phase": "生成复习卡片..."}

        from schemas import NoteOut
        note_data = NoteOut.model_validate(note, from_attributes=True).model_dump()
        _generate_tasks[task_id] = {"status": "done", "progress": 1.0, "note": note_data}
    except Exception as e:
        _generate_tasks[task_id] = {"status": "error", "error": str(e), "progress": 0}
    finally:
        db.close()


@router.post("/{video_id}/generate-async")
def start_generate(video_id: int, data: GenerateIn | None = None, db: Session = Depends(get_db), x_bili_token: str | None = Header(None)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")

    text = data.text if data and data.text else ""
    task_id = str(uuid.uuid4())
    _generate_tasks[task_id] = {"status": "queued", "progress": 0, "phase": "等待中..."}

    t = threading.Thread(target=_run_generate, args=(video_id, task_id, text, x_bili_token), daemon=True)
    t.start()

    return {"task_id": task_id, "note": None}


@router.get("/{video_id}/generate/status")
def get_generate_status(video_id: int, task_id: str):
    task = _generate_tasks.get(task_id)
    if not task:
        return {"status": "not_found", "progress": 0}

    if task["status"] == "done":
        return task
    return {"status": task["status"], "progress": task["progress"], "phase": task.get("phase", ""), "error": task.get("error", "")}


# ── Download helpers ──────────────────────────────────────────

def _build_note_data(v: Video, note: Note | None, db: Session) -> dict:
    points = []
    cards = []
    if note:
        try:
            points = json.loads(note.key_points) if isinstance(note.key_points, str) else note.key_points
        except Exception:
            points = []
        cards = db.query(ReviewCard).filter(ReviewCard.note_id == note.id).all()
    return {
        "title": v.title,
        "source": v.source,
        "bvid": v.bvid,
        "duration": v.duration,
        "subtitle_text": note.subtitle_text if note else "",
        "cleaned_subtitle": note.cleaned_subtitle if note else "",
        "translated_subtitle": note.translated_subtitle if note else "",
        "summary": note.summary if note else "",
        "key_points": points,
        "cards": cards,
    }


def _download_md(data: dict) -> PlainTextResponse:
    platform = "YouTube" if data["source"] == "youtube" else "B站"
    lines = [
        f"# {data['title']}",
        "",
        f"- 来源: {platform}",
        f"- ID: {data['bvid']}",
        f"- 时长: {data['duration']}秒",
        "",
        "---",
        "",
    ]
    if data["subtitle_text"]:
        lines += ["## 📝 字幕原文", "", data["subtitle_text"], ""]
    if data["cleaned_subtitle"]:
        lines += ["## ✨ 字幕精校版", "", data["cleaned_subtitle"], ""]
    if data["subtitle_text"]:
        lines += ["---", ""]
    if data["summary"]:
        lines += ["## 📖 AI 笔记总结", "", data["summary"], ""]
    if data["key_points"]:
        lines += ["", "## 🎯 核心知识点", ""]
        for i, p in enumerate(data["key_points"], 1):
            lines += [f"{i}. {p}", ""]
    if data["cards"]:
        lines += ["", "## 💡 复习卡片", ""]
        for i, c in enumerate(data["cards"], 1):
            lines += [f"### 卡片 {i}", "", f"**Q:** {c.question}", "", f"**A:** {c.answer}", ""]

    content = "\n".join(lines)
    filename = f"BilLeaRN_{data['bvid']}.md"
    return PlainTextResponse(content, headers={
        "Content-Disposition": f'attachment; filename="{filename}"',
        "Content-Type": "text/markdown; charset=utf-8",
    })


def _pdf_setup_cjk(pdf):
    for path in [r"C:\Windows\Fonts\simhei.ttf"]:
        if os.path.exists(path):
            pdf.add_font("CJK", "", path)
            pdf.add_font("CJK", "B", path)
            return True
    return False


def _download_pdf(data: dict) -> Response:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    FAM = "CJK" if _pdf_setup_cjk(pdf) else "Helvetica"

    def mc(w, h, txt, fam=FAM, size=None, style=""):
        if size: pdf.set_font(fam, style, size)
        pdf.multi_cell(w, h, txt, new_x="LMARGIN")

    mc(0, 8, data["title"], size=16, style="B")
    pdf.ln(4)
    platform = "YouTube" if data["source"] == "youtube" else "B站"
    mc(0, 6, f"来源: {platform}  |  ID: {data['bvid']}  |  时长: {data['duration']}秒", size=10)
    pdf.ln(6)

    if data["subtitle_text"]:
        mc(0, 7, "字幕原文", size=13, style="B")
        mc(0, 5, data["subtitle_text"], size=9)
        pdf.ln(4)

    if data["cleaned_subtitle"]:
        mc(0, 7, "字幕精校版", size=13, style="B")
        mc(0, 5, data["cleaned_subtitle"], size=9)
        pdf.ln(4)

    if data.get("translated_subtitle"):
        mc(0, 7, "双语翻译", size=13, style="B")
        mc(0, 5, data["translated_subtitle"], size=9)
        pdf.ln(4)

    if data["summary"]:
        mc(0, 7, "AI 笔记总结", size=13, style="B")
        mc(0, 6, data["summary"], size=10)
        pdf.ln(4)

    if data["key_points"]:
        mc(0, 7, "核心知识点", size=13, style="B")
        for i, p in enumerate(data["key_points"], 1):
            mc(0, 6, f"{i}. {p}", size=10)
            pdf.ln(2)

    if data["cards"]:
        mc(0, 7, "复习卡片", size=13, style="B")
        for i, c in enumerate(data["cards"], 1):
            mc(0, 6, f"卡片 {i}", size=10, style="B")
            mc(0, 6, f"Q: {c.question}", size=10)
            mc(0, 6, f"A: {c.answer}", size=10)
            pdf.ln(3)

    pdf_bytes = bytes(pdf.output())
    filename = f"BilLeaRN_{data['bvid']}.pdf"
    return Response(pdf_bytes, media_type="application/pdf", headers={
        "Content-Disposition": f'attachment; filename="{filename}"',
    })


def _download_docx(data: dict) -> Response:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    doc.add_heading(data["title"], level=1)
    platform = "YouTube" if data["source"] == "youtube" else "B站"
    p = doc.add_paragraph(f"来源: {platform}  |  ID: {data['bvid']}  |  时长: {data['duration']}秒")
    p.runs[0].font.size = Pt(9)
    doc.add_paragraph("")

    if data["subtitle_text"]:
        doc.add_heading("字幕原文", level=2)
        doc.add_paragraph(data["subtitle_text"])

    if data["cleaned_subtitle"]:
        doc.add_heading("字幕精校版", level=2)
        doc.add_paragraph(data["cleaned_subtitle"])

    if data["summary"]:
        doc.add_heading("AI 笔记总结", level=2)
        doc.add_paragraph(data["summary"])

    if data["key_points"]:
        doc.add_heading("核心知识点", level=2)
        for i, p in enumerate(data["key_points"], 1):
            doc.add_paragraph(f"{i}. {p}")

    if data["cards"]:
        doc.add_heading("复习卡片", level=2)
        for i, c in enumerate(data["cards"], 1):
            doc.add_heading(f"卡片 {i}", level=3)
            doc.add_paragraph(f"Q: {c.question}")
            doc.add_paragraph(f"A: {c.answer}")

    buf = io.BytesIO()
    doc.save(buf)
    filename = f"BilLeaRN_{data['bvid']}.docx"
    return Response(buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
        "Content-Disposition": f'attachment; filename="{filename}"',
    })


# ── Download endpoints ──────────────────────────────────────────

@router.get("/{video_id}/download")
def download_note_md(video_id: int, db: Session = Depends(get_db)):
    return _download_note(video_id, "md", db)


@router.get("/{video_id}/download/md")
def download_note_md_explicit(video_id: int, db: Session = Depends(get_db)):
    return _download_note(video_id, "md", db)


@router.get("/{video_id}/download/pdf")
def download_note_pdf(video_id: int, db: Session = Depends(get_db)):
    return _download_note(video_id, "pdf", db)


@router.get("/{video_id}/download/docx")
def download_note_docx(video_id: int, db: Session = Depends(get_db)):
    return _download_note(video_id, "docx", db)


def _download_note(video_id: int, fmt: str, db: Session):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")
    note = db.query(Note).filter(Note.video_id == video_id).first()
    data = _build_note_data(v, note, db)

    if fmt == "pdf":
        return _download_pdf(data)
    elif fmt == "docx":
        return _download_docx(data)
    return _download_md(data)


@router.get("/{video_id}/audio")
def download_audio(video_id: int, db: Session = Depends(get_db)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(404, "Video not found")
    from services.transcribe import get_audio_path
    path = get_audio_path(v.url, v.bvid)
    if not path or not os.path.exists(path):
        raise HTTPException(500, "音频下载失败")
    filename = f"{v.bvid}.m4a"
    return FileResponse(path, media_type="audio/mp4", filename=filename,
                        headers={"Content-Disposition": f'attachment; filename="{filename}"'})
